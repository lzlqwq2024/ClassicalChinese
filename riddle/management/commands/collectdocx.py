from django.core.management.base import BaseCommand, CommandError
from riddle.models import Article
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

class Command(BaseCommand):
    help = 'Collect the docx for installing'

    def handle(self, *args, **options):
        try:
            articlelist = Article.objects.all()
            length = str(len(articlelist))
            path = os.path.abspath(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
            pk = 0
            for article in articlelist:
                if os.path.exists(path+'\\static\\riddle\\docx\\'+str(article.number)+article.name+'.docx'):
                    continue
                wordfile = Document()

                paragraph1 = wordfile.add_paragraph()
                BigTitle = paragraph1.add_run(str(article.number)+"."+" "+article.name)
                BigTitle.font.name = '楷体'
                bt = BigTitle._element.rPr.rFonts
                bt.set(qn('w:eastAsia'), '楷体')
                BigTitle.font.size = Pt(18)
                BigTitle.font.bold = True
                BigTitle.add_break()

                paragraph2 = wordfile.add_paragraph()
                SmallTitle1 = paragraph2.add_run('文言文原文')
                SmallTitle1.font.name = '楷体'
                st1 = SmallTitle1._element.rPr.rFonts
                st1.set(qn('w:eastAsia'), '楷体')
                SmallTitle1.font.size = Pt(16)
                SmallTitle1.font.bold = True

                paragraph3 = wordfile.add_paragraph()
                original_text = paragraph3.add_run(article.text)
                original_text.font.name = '楷体'
                ot = original_text._element.rPr.rFonts
                ot.set(qn('w:eastAsia'), '楷体')
                original_text.font.size = Pt(14)
                paragraph3.paragraph_format.space_before = Pt(15)
                paragraph3.paragraph_format.line_spacing = 2
                original_text.add_break()

                paragraph4 = wordfile.add_paragraph()
                SmallTitle2 = paragraph4.add_run('文言文大意')
                SmallTitle2.font.name = '楷体'
                st2 = SmallTitle2._element.rPr.rFonts
                st2.set(qn('w:eastAsia'), '楷体')
                SmallTitle2.font.size = Pt(16)
                SmallTitle2.font.bold = True

                paragraph5 = wordfile.add_paragraph()
                translate = paragraph5.add_run(article.translate)
                translate.font.name = '楷体'
                ts = translate._element.rPr.rFonts
                ts.set(qn('w:eastAsia'), '楷体')
                translate.font.size = Pt(14)
                paragraph5.paragraph_format.space_before = Pt(15)
                paragraph5.paragraph_format.line_spacing = 2
                translate.add_break()

                paragraph6 = wordfile.add_paragraph()
                SmallTitle3 = paragraph6.add_run('文言文注释')
                SmallTitle3.font.name = '楷体'
                st3 = SmallTitle3._element.rPr.rFonts
                st3.set(qn('w:eastAsia'), '楷体')
                SmallTitle3.font.size = Pt(16)
                SmallTitle3.font.bold = True

                paragraph7 = wordfile.add_paragraph()
                annotation = paragraph7.add_run(article.annotation)
                annotation.font.name = '楷体'
                at = annotation._element.rPr.rFonts
                at.set(qn('w:eastAsia'), '楷体')
                annotation.font.size = Pt(14)
                paragraph7.paragraph_format.space_before = Pt(15)
                paragraph7.paragraph_format.line_spacing = 2
                annotation.add_break()

                wordfile.save(path+'\\static\\riddle\\docx\\'+str(article.number)+article.name+'.docx')
                pk = pk+1
            self.stdout.write(self.style.SUCCESS('Successfully collect %s new docx' % pk)) 
        except:
            raise CommandError('Collect docx failed')